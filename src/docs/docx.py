import docx
from func.translation import translate
from func.async_translation import get_batch_translation
from globals import counter

from tqdm import tqdm


class Docx():
    def __init__(self, path):
        self.path = path
        self.rawdoc = docx.Document(self.path)


    def rebuild(self, name="translated.docx", batch=5):
        SKIP = '1234567890!"№;%:?*),."«»'
        n = 0
        tasks = []
        for t in self.rawdoc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            if r.text and r.text not in SKIP and r.text.replace("\n",""):
                                n+=1
        for p in self.rawdoc.paragraphs:
            for r in p.runs:
                if r.text and r.text not in SKIP:
                    n+=1
        
        with tqdm(total=n) as pbar:
            for i, paragraph in enumerate(self.rawdoc.paragraphs):
                if len(tasks) == batch:
                    results = get_batch_translation(tasks)
                    for result in results:
                        for i,run in enumerate(result[0]):
                            run.text = result[1][i]
                            pbar.update(1)
                    print(counter.count)
                    tasks = []
        
                else:
                    runs = [run for run in paragraph.runs if run.text and run.text not in SKIP]
                    text = [run.text for run in paragraph.runs if run.text and run.text not in SKIP]
                    if text:
                        tasks.append([runs,text])
            if tasks:
                results = get_batch_translation(tasks)
                for result in results:
                    for i,run in enumerate(result[0]):
                        run.text = result[1]
                        pbar.update(1)
                tasks = []

            total = counter.count
            print(total)

                    
      
            for n, table in enumerate(self.rawdoc.tables):
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        for k, paragraph in enumerate(cell.paragraphs):
                            runs = [run for run in paragraph.runs if run.text and run.text not in SKIP and run.text.replace("\n","")]
                            text = [run.text for run in paragraph.runs if run.text and run.text not in SKIP and run.text.replace("\n","")]
                            text_translated = translate(text)
                            for i,run in enumerate(runs):
                                run.text = text_translated[i]
                                pbar.update(1)

        self.rawdoc.save(name)
